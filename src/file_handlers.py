"""
Модуль для чтения различных форматов файлов
"""

import os
import re
from abc import ABC, abstractmethod
from typing import Optional, Dict, Any
from pathlib import Path
import structlog

# Импорты для обработки файлов
try:
    import docx
except ImportError:
    docx = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

from .utils import get_file_extension, format_file_size

logger = structlog.get_logger(__name__)


class FileHandlerError(Exception):
    """Исключение для ошибок обработки файлов"""
    pass


class BaseFileHandler(ABC):
    """Базовый класс для обработчиков файлов"""
    
    def __init__(self):
        self.supported_extensions = []
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """
        Извлечение текста из файла
        
        Args:
            file_path: Путь к файлу
        
        Returns:
            Извлеченный текст
        """
        pass
    
    def validate_file(self, file_path: str) -> bool:
        """
        Валидация файла
        
        Args:
            file_path: Путь к файлу
        
        Returns:
            True если файл валиден
        """
        if not os.path.exists(file_path):
            return False
        
        if not os.path.isfile(file_path):
            return False
        
        extension = get_file_extension(file_path)
        return extension in self.supported_extensions
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Получение информации о файле
        
        Args:
            file_path: Путь к файлу
        
        Returns:
            Словарь с информацией о файле
        """
        path = Path(file_path)
        stat = path.stat()
        
        return {
            "name": path.name,
            "size": stat.st_size,
            "size_formatted": format_file_size(stat.st_size),
            "extension": get_file_extension(file_path),
            "modified": stat.st_mtime
        }


class TextFileHandler(BaseFileHandler):
    """Обработчик для обычных текстовых файлов"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.txt']
    
    def extract_text(self, file_path: str) -> str:
        """Извлечение текста из .txt файла"""
        if not self.validate_file(file_path):
            raise FileHandlerError(f"Некорректный .txt файл: {file_path}")
        
        # Попытка определить кодировку
        encodings = ['utf-8', 'utf-8-sig', 'cp1251', 'cp866', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                
                logger.debug(
                    "Текстовый файл успешно прочитан",
                    file_path=file_path,
                    encoding=encoding,
                    length=len(text)
                )
                return text
                
            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise FileHandlerError(f"Ошибка чтения файла {file_path}: {e}")
        
        raise FileHandlerError(f"Не удалось определить кодировку файла: {file_path}")


class MarkdownFileHandler(BaseFileHandler):
    """Обработчик для Markdown файлов"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.md', '.markdown']
    
    def extract_text(self, file_path: str) -> str:
        """Извлечение текста из .md файла с базовой обработкой разметки"""
        if not self.validate_file(file_path):
            raise FileHandlerError(f"Некорректный .md файл: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Базовая обработка Markdown разметки
            text = self._process_markdown(content)
            
            logger.debug(
                "Markdown файл успешно прочитан",
                file_path=file_path,
                original_length=len(content),
                processed_length=len(text)
            )
            return text
            
        except Exception as e:
            raise FileHandlerError(f"Ошибка чтения Markdown файла {file_path}: {e}")
    
    def _process_markdown(self, content: str) -> str:
        """Базовая обработка Markdown разметки"""
        # Удаляем заголовки (оставляем только текст)
        content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
        
        # Удаляем ссылки, оставляем только текст
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        
        # Удаляем жирный и курсивный текст (оставляем содержимое)
        content = re.sub(r'\*\*([^\*]+)\*\*', r'\1', content)
        content = re.sub(r'\*([^\*]+)\*', r'\1', content)
        content = re.sub(r'__([^_]+)__', r'\1', content)
        content = re.sub(r'_([^_]+)_', r'\1', content)
        
        # Удаляем код блоки
        content = re.sub(r'```[^`]*```', '', content, flags=re.DOTALL)
        content = re.sub(r'`([^`]+)`', r'\1', content)
        
        # Удаляем списки (оставляем только текст)
        content = re.sub(r'^\s*[-*+]\s+', '', content, flags=re.MULTILINE)
        content = re.sub(r'^\s*\d+\.\s+', '', content, flags=re.MULTILINE)
        
        # Очищаем лишние пробелы и переносы
        content = re.sub(r'\n\s*\n', '\n\n', content)
        content = content.strip()
        
        return content


class DocxFileHandler(BaseFileHandler):
    """Обработчик для .docx файлов"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']
        
        if docx is None:
            logger.warning("python-docx не установлен, .docx файлы не поддерживаются")
    
    def extract_text(self, file_path: str) -> str:
        """Извлечение текста из .docx файла"""
        if docx is None:
            raise FileHandlerError("python-docx не установлен")
        
        if not self.validate_file(file_path):
            raise FileHandlerError(f"Некорректный .docx файл: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Извлекаем текст из всех параграфов
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            paragraphs.append(text)
            
            full_text = '\n\n'.join(paragraphs)
            
            logger.debug(
                "DOCX файл успешно прочитан",
                file_path=file_path,
                paragraphs_count=len(paragraphs),
                length=len(full_text)
            )
            return full_text
            
        except Exception as e:
            raise FileHandlerError(f"Ошибка чтения DOCX файла {file_path}: {e}")


class PdfFileHandler(BaseFileHandler):
    """Обработчик для .pdf файлов"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
        
        if PyPDF2 is None and pdfplumber is None:
            logger.warning("PyPDF2 и pdfplumber не установлены, .pdf файлы не поддерживаются")
    
    def extract_text(self, file_path: str) -> str:
        """Извлечение текста из .pdf файла"""
        if PyPDF2 is None and pdfplumber is None:
            raise FileHandlerError("PyPDF2 или pdfplumber должны быть установлены")
        
        if not self.validate_file(file_path):
            raise FileHandlerError(f"Некорректный .pdf файл: {file_path}")
        
        # Пробуем сначала pdfplumber (обычно лучше извлекает текст)
        if pdfplumber is not None:
            try:
                return self._extract_with_pdfplumber(file_path)
            except Exception as e:
                logger.warning(f"pdfplumber не смог обработать файл: {e}")
        
        # Если pdfplumber не сработал, пробуем PyPDF2
        if PyPDF2 is not None:
            try:
                return self._extract_with_pypdf2(file_path)
            except Exception as e:
                logger.warning(f"PyPDF2 не смог обработать файл: {e}")
        
        raise FileHandlerError(f"Не удалось извлечь текст из PDF файла: {file_path}")
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Извлечение текста с помощью pdfplumber"""
        pages_text = []
        
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text.strip())
                except Exception as e:
                    logger.warning(f"Ошибка извлечения текста со страницы {page_num}: {e}")
        
        full_text = '\n\n'.join(pages_text)
        
        logger.debug(
            "PDF файл успешно прочитан (pdfplumber)",
            file_path=file_path,
            pages_count=len(pages_text),
            length=len(full_text)
        )
        return full_text
    
    def _extract_with_pypdf2(self, file_path: str) -> str:
        """Извлечение текста с помощью PyPDF2"""
        pages_text = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text.strip())
                except Exception as e:
                    logger.warning(f"Ошибка извлечения текста со страницы {page_num}: {e}")
        
        full_text = '\n\n'.join(pages_text)
        
        logger.debug(
            "PDF файл успешно прочитан (PyPDF2)",
            file_path=file_path,
            pages_count=len(pages_text),
            length=len(full_text)
        )
        return full_text


class FileHandlerFactory:
    """Фабрика для создания обработчиков файлов"""
    
    def __init__(self):
        self.handlers = {
            '.txt': TextFileHandler(),
            '.md': MarkdownFileHandler(),
            '.markdown': MarkdownFileHandler(),
            '.docx': DocxFileHandler(),
            '.pdf': PdfFileHandler()
        }
    
    def get_handler(self, file_path: str) -> Optional[BaseFileHandler]:
        """
        Получение обработчика для файла
        
        Args:
            file_path: Путь к файлу
        
        Returns:
            Обработчик файла или None
        """
        extension = get_file_extension(file_path)
        return self.handlers.get(extension)
    
    def is_supported(self, file_path: str) -> bool:
        """
        Проверка поддержки формата файла
        
        Args:
            file_path: Путь к файлу
        
        Returns:
            True если формат поддерживается
        """
        return self.get_handler(file_path) is not None
    
    def get_supported_extensions(self) -> list:
        """
        Получение списка поддерживаемых расширений
        
        Returns:
            Список расширений
        """
        return list(self.handlers.keys())


# Глобальная фабрика обработчиков
_file_factory: Optional[FileHandlerFactory] = None


def get_file_factory() -> FileHandlerFactory:
    """
    Получение глобальной фабрики обработчиков
    
    Returns:
        Экземпляр FileHandlerFactory
    """
    global _file_factory
    if _file_factory is None:
        _file_factory = FileHandlerFactory()
    return _file_factory


def extract_text_from_file(file_path: str) -> str:
    """
    Удобная функция для извлечения текста из файла
    
    Args:
        file_path: Путь к файлу
    
    Returns:
        Извлеченный текст
    
    Raises:
        FileHandlerError: Если файл не поддерживается или произошла ошибка
    """
    factory = get_file_factory()
    handler = factory.get_handler(file_path)
    
    if handler is None:
        supported = factory.get_supported_extensions()
        raise FileHandlerError(
            f"Неподдерживаемый формат файла: {get_file_extension(file_path)}. "
            f"Поддерживаемые форматы: {', '.join(supported)}"
        )
    
    return handler.extract_text(file_path)


def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Получение информации о файле
    
    Args:
        file_path: Путь к файлу
    
    Returns:
        Словарь с информацией о файле
    """
    factory = get_file_factory()
    handler = factory.get_handler(file_path)
    
    if handler is None:
        # Базовая информация для неподдерживаемых файлов
        path = Path(file_path)
        if path.exists():
            stat = path.stat()
            return {
                "name": path.name,
                "size": stat.st_size,
                "size_formatted": format_file_size(stat.st_size),
                "extension": get_file_extension(file_path),
                "modified": stat.st_mtime,
                "supported": False
            }
        else:
            raise FileHandlerError(f"Файл не найден: {file_path}")
    
    info = handler.get_file_info(file_path)
    info["supported"] = True
    return info


def validate_input_file(file_path: str) -> bool:
    """
    Валидация входного файла
    
    Args:
        file_path: Путь к файлу
    
    Returns:
        True если файл валиден для обработки
    """
    try:
        factory = get_file_factory()
        handler = factory.get_handler(file_path)
        
        if handler is None:
            return False
        
        return handler.validate_file(file_path)
        
    except Exception:
        return False
